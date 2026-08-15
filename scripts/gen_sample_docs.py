"""生成示例 PDF / Word 知识库文档（一次性工具，可按需重新生成）。

用法：PYTHONPATH=. .venv/bin/python scripts/gen_sample_docs.py
说明：PDF 中文内容依赖系统中文字体，找不到时自动退回英文内容。
"""
import os

PROJECT_ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
KNOWLEDGE_DIR = os.path.join(PROJECT_ROOT, "docs", "knowledge")

PDF_CONTENT = [
    ("DocMind 部署与运维指南", 1),
    ("运行环境要求", 2),
    ("DocMind 需要 Python 3.10 及以上版本。依赖安装使用 pip install -r requirements.txt，"
     "国内网络建议加阿里云镜像源：-i https://mirrors.aliyun.com/pypi/simple/。"
     "百炼 API Key 需要写入项目根目录的 .env 文件，文件不会被提交到 Git。", 0),
    ("启动方式", 2),
    ("命令行模式：python -m docmind.cli，适合调试。"
     "Web 界面模式：python -m docmind.app，默认监听 127.0.0.1 的 7860 端口。"
     "修改端口可在 app.py 的 launch 调用中传入 server_port 参数。", 0),
    ("知识库管理", 2),
    ("知识库文档存放在 docs/knowledge 目录，支持 md、txt、pdf、docx 四种格式。"
     "新增或删除文档后需要重启应用以重建索引。单个文档解析失败只告警跳过，不影响其余文档。", 0),
    ("常见问题", 2),
    ("端口被占用：先执行 lsof -i:7860 找到占用进程再结束它。"
     "API 限流：系统内置退避重试，持续失败请检查百炼控制台额度。"
     "MCP 工具不可用：确认 .venv 解释器路径正确，天气服务由主进程自动拉起。", 0),
]

DOCX_CONTENT = {
    "title": "DocMind 常见问题排查手册",
    "paragraphs": [
        "本手册汇总 DocMind 运行中最常见的故障现象、原因与解决办法，按模块分类。",
    ],
    "sections": [
        ("API 调用类故障", [
            "现象：回答中提示模型调用失败。原因通常是 API Key 无效、欠费或限流。"
            "解决办法：核对 .env 中的 DASHSCOPE_API_KEY，到百炼控制台确认额度；"
            "系统已内置三次退避重试，偶发失败会自动恢复。",
            "现象：启动时报未配置 DASHSCOPE_API_KEY。原因是 .env 文件缺失或字段名拼错。"
            "解决办法：复制 .env.example 为 .env 并填入真实 Key。",
        ]),
        ("知识库类故障", [
            "现象：新增文档后问答没有引用新内容。原因是知识库索引在启动时构建。"
            "解决办法：重启应用，观察启动日志中的切片数量是否增加。",
            "现象：某个 PDF 或 Word 文档没有被加载。原因可能是文件损坏或加密。"
            "解决办法：查看启动日志中的解析告警，重新导出该文档。",
        ]),
        ("工具调用类故障", [
            "现象：天气查询不可用。原因是 MCP Server 子进程启动失败。"
            "解决办法：手动运行 python mcp_servers/weather_server.py 查看报错。",
        ]),
    ],
    "table": {
        "headers": ["故障现象", "可能原因", "处理动作"],
        "rows": [
            ["回答一直显示思考中", "API 超时", "等待自动重试或检查网络"],
            ["检索结果不相关", "问题过于口语化", "换用文档中的关键词重新提问"],
            ["界面样式错乱", "浏览器缓存旧样式", "强制刷新页面"],
        ],
    },
}

# 常见中文字体路径（macOS）
CJK_FONT_CANDIDATES = [
    "/System/Library/Fonts/PingFang.ttc",
    "/System/Library/Fonts/STHeiti Light.ttc",
    "/System/Library/Fonts/Hiragino Sans GB.ttc",
]


def gen_pdf(path: str):
    from fpdf import FPDF

    pdf = FPDF()
    pdf.add_page()
    cjk = next((f for f in CJK_FONT_CANDIDATES if os.path.exists(f)), None)
    use_cjk = False
    if cjk:
        try:
            pdf.add_font("CJK", "", cjk)
            pdf.add_font("CJK", "B", cjk)
            use_cjk = True
        except Exception as e:  # noqa: BLE001 - 字体加载失败退回英文
            print(f"[提示] 中文字体加载失败，PDF 退回英文内容: {e}")
    if not use_cjk:
        pdf.set_font("Helvetica", size=11)
        pdf.multi_cell(0, 7, "CJK font unavailable; see docs/knowledge markdown files for Chinese content.")

    for text, level in PDF_CONTENT:
        if not use_cjk:
            continue
        if level == 1:
            pdf.set_font("CJK", "B", 18)
            pdf.multi_cell(0, 11, text)
            pdf.ln(3)
        elif level == 2:
            pdf.set_font("CJK", "B", 14)
            pdf.ln(2)
            pdf.multi_cell(0, 9, text)
            pdf.ln(1)
        else:
            pdf.set_font("CJK", "", 11)
            pdf.multi_cell(0, 7, text)
    pdf.output(path)
    print(f"已生成 {path}（{'中文' if use_cjk else '英文回退'}）")


def gen_docx(path: str):
    import docx

    d = docx.Document()
    d.add_heading(DOCX_CONTENT["title"], level=0)
    for p in DOCX_CONTENT["paragraphs"]:
        d.add_paragraph(p)
    for heading, paras in DOCX_CONTENT["sections"]:
        d.add_heading(heading, level=1)
        for p in paras:
            d.add_paragraph(p)
    t = DOCX_CONTENT["table"]
    table = d.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(t["headers"]):
        hdr[i].text = h
    for row in t["rows"]:
        cells = table.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = v
    d.save(path)
    print(f"已生成 {path}")


if __name__ == "__main__":
    os.makedirs(KNOWLEDGE_DIR, exist_ok=True)
    gen_pdf(os.path.join(KNOWLEDGE_DIR, "部署与运维指南.pdf"))
    gen_docx(os.path.join(KNOWLEDGE_DIR, "常见问题排查手册.docx"))
