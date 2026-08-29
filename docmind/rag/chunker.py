"""文档加载与切片：把 docs/knowledge 下的文档切成可检索的片段。

格式支持（面试可讲）：
- .md / .txt：直接读取
- .pdf：pypdf 逐页提取（纯 Python，无系统依赖）
- .docx：python-docx 提取正文段落 + 表格
- .xlsx：openpyxl 按 Sheet 提取，行数据管道符拼接（与 docx 表格策略一致）
- .png/.jpg/.jpeg/.webp：百炼多模态 OCR 抽取图中文字（结果磁盘缓存，避免重复调 API）
- 单个文件解析失败只告警跳过，不影响其余文档建库

切片策略：
- Markdown 文档优先按标题（# ~ ####）切分成语义完整的段落，再合并小段、
  拆分超长段，让每个切片尽量是一个完整的 QA 或主题
- 纯文本/PDF/Word 退回滑窗切片：CHUNK_SIZE 窗口 + CHUNK_OVERLAP 重叠，优先换行处断开
"""
import base64
import hashlib
import logging
import os
import pathlib
import re

from docmind import config

logger = logging.getLogger(__name__)

SUPPORTED_EXTS = {".md", ".txt", ".pdf", ".docx", ".xlsx", ".csv", ".json", ".png", ".jpg", ".jpeg", ".webp"}
IMAGE_EXTS = {".png", ".jpg", ".jpeg", ".webp"}
# 解析器资源上限：恶意构造的文档(zip bomb/万页 PDF)会拖垮建库进程
_MAX_EXTRACT_CHARS = int(os.getenv("MAX_EXTRACT_CHARS", str(4 * 1024 * 1024)))  # 单文档提取文本上限 4MB
_MAX_PDF_PAGES = int(os.getenv("MAX_PDF_PAGES", "500"))                          # 单 PDF 最大解析页数
_MAX_XLSX_ROWS = int(os.getenv("MAX_XLSX_ROWS", "200000"))                       # 单 xlsx 最大行数
OCR_CACHE_DIR = os.path.join(config.PROJECT_ROOT, "data", "ocr_cache")
_HEADING_RE = re.compile(r"(?=^#{1,4}\s)", re.MULTILINE)  # 零宽断言：切分但保留标题文本


# ---------------- 格式提取器 ----------------
# 解析器资源上限:防恶意构造文件耗尽内存(zip bomb / 超大页)
_MAX_DOCX_UNCOMPRESSED = 200 * 1024 * 1024   # 解压后总大小上限 200MB


def _extract_pdf_pages(path: str) -> list[tuple[int, str]]:
    """pypdf 逐页提取文本，返回 [(页号, 文本)]（页号从 1 起），供切片携带页码元数据"""
    from pypdf import PdfReader

    reader = PdfReader(path)
    if len(reader.pages) > _MAX_PDF_PAGES:
        raise ValueError(f"PDF 页数超过上限 {_MAX_PDF_PAGES}")
    pages = []
    for i, page in enumerate(reader.pages, 1):
        if i > _MAX_PDF_PAGES:
            logger.warning(f"PDF 超过 {_MAX_PDF_PAGES} 页，仅解析前 {_MAX_PDF_PAGES} 页: {os.path.basename(path)}")
            break
        text = (page.extract_text() or "").strip()
        if text:
            pages.append((i, text))
    return pages


def _extract_pdf(path: str) -> str:
    """pypdf 逐页提取文本，页间用空行分隔"""
    return "\n\n".join(t for _, t in _extract_pdf_pages(path))


def _extract_docx(path: str) -> str:
    """python-docx 提取正文段落与表格（表格按行拼接）"""
    import zipfile

    with zipfile.ZipFile(path) as _zf:   # docx 本质是 zip:防解压炸弹
        total = sum(i.file_size for i in _zf.infolist())
        if total > _MAX_DOCX_UNCOMPRESSED:
            raise ValueError(f"docx 解压后超过上限 {_MAX_DOCX_UNCOMPRESSED // (1024 * 1024)}MB")

    import docx

    d = docx.Document(path)
    parts = [p.text.strip() for p in d.paragraphs if p.text.strip()]
    for table in d.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n\n".join(parts)


def _extract_xlsx(path: str) -> str:
    """openpyxl 逐 Sheet 提取：行内单元格管道符拼接；大 Sheet 预分组，
    每组带 [Sheet: 名] 标记并重复表头——切片后每个分片都保有完整列语义"""
    import openpyxl

    wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
    parts = []
    try:
        for ws in wb.worksheets:
            rows = []
            for row in ws.iter_rows(values_only=True):
                cells = [str(c).strip() for c in row if c is not None and str(c).strip()]
                if cells:
                    rows.append(" | ".join(cells))
            if not rows:
                continue
            marker = f"[Sheet: {ws.title}]"
            header = rows[0]
            groups, cur, used = [], [marker, header], len(marker) + len(header)
            for r in rows[1:]:
                if used + len(r) + 1 > config.CHUNK_SIZE and len(cur) > 2:
                    groups.append("\n".join(cur))
                    cur, used = [marker, header], len(marker) + len(header)
                cur.append(r)
                used += len(r) + 1
            groups.append("\n".join(cur))
            parts.extend(groups)
    finally:
        wb.close()
    return "\n\n".join(parts)


def _ocr_image(path: str) -> str:
    """图片 OCR：百炼多模态抽取图中文字。按「文件名+大小+mtime」磁盘缓存，
    索引重建时不重复调 API；OCR 文本为空视为失败（抛错由建库流程告警跳过）"""
    st = os.stat(path)
    fp = hashlib.sha256(
        f"{os.path.basename(path)}:{st.st_size}:{st.st_mtime_ns}".encode()
    ).hexdigest()[:16]
    os.makedirs(OCR_CACHE_DIR, exist_ok=True)
    cache_path = pathlib.Path(OCR_CACHE_DIR, fp + ".txt")
    if cache_path.is_file():
        return cache_path.read_text(encoding="utf-8")

    import requests

    with open(path, "rb") as f:
        b64 = base64.b64encode(f.read()).decode()
    ext = os.path.splitext(path)[1].lstrip(".").lower()
    mime = "jpeg" if ext == "jpg" else ext
    resp = requests.post(
        config.DASHSCOPE_BASE_URL.rstrip("/") + "/chat/completions",
        headers={"Authorization": f"Bearer {config.DASHSCOPE_API_KEY}"},
        json={
            "model": config.OCR_MODEL,
            "messages": [{"role": "user", "content": [
                {"type": "image_url",
                 "image_url": {"url": f"data:image/{mime};base64,{b64}"}},
                {"type": "text", "text": "请读取图片中的全部文字并原样输出，不要添加解释。"},
            ]}],
        },
        timeout=60,
    )
    resp.raise_for_status()
    text = (resp.json()["choices"][0]["message"]["content"] or "").strip()
    if not text:
        raise ValueError("OCR 未识别到文字")
    cache_path.write_text(text, encoding="utf-8")
    return text


# 二进制格式的提取器映射（懒导入，不用时不加载库）
_EXTRACTORS = {
    ".pdf": _extract_pdf,
    ".docx": _extract_docx,
    ".xlsx": _extract_xlsx,
    **{ext: _ocr_image for ext in IMAGE_EXTS},
}


def load_documents(knowledge_dir: str | None = None) -> list[dict]:
    """读取知识库目录，返回 [{source, text}]；单文件失败跳过不阻断"""
    root = knowledge_dir or config.KNOWLEDGE_DIR
    docs = []
    if not os.path.isdir(root):
        return docs
    for name in sorted(os.listdir(root)):
        ext = os.path.splitext(name)[1].lower()
        if ext not in SUPPORTED_EXTS:
            continue
        path = os.path.join(root, name)
        try:
            if ext in _EXTRACTORS:
                text = _EXTRACTORS[ext](path)
            else:
                with open(path, encoding="utf-8") as f:
                    text = f.read()
        except Exception as e:  # noqa: BLE001 - 坏文件不能阻断整个建库流程
            logger.warning(f"文档解析失败，已跳过 {name}: {e}")
            continue
        text = (text or "")[:_MAX_EXTRACT_CHARS]
        if text.strip():
            docs.append({"source": name, "text": text.strip()})
    return docs


def _window_split(text: str) -> list[str]:
    """滑窗切片：优先沿句子边界断开。

    原实现只在窗口后半段找换行——密集中文长段落（无空行、句号密集）
    前半段永远找不到断点而被硬切。现改为全窗搜索句末标点
    （。！？；\\n 优先级递减），并要求断点不早于窗口 1/4 处（防碎片）。"""
    size, overlap = config.CHUNK_SIZE, config.CHUNK_OVERLAP
    min_break = max(1, size // 4)
    chunks: list[str] = []
    start = 0
    while start < len(text):
        end = min(start + size, len(text))
        if end < len(text):
            window = text[start:end]
            for punct in ("。", "！", "？", "；", "\n"):
                pos = window.rfind(punct, min_break - 1)
                if pos != -1:
                    end = start + pos + len(punct)
                    break
            else:
                # 无句末标点：退而求其次找空格（英文），再退硬切
                sp = window.rfind(" ", min_break - 1)
                if sp > 0:
                    end = start + sp + 1
        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)
        if end >= len(text):
            break
        start = max(end - overlap, start + 1)
    return chunks


def _is_table_row(line: str) -> bool:
    """表格行判定：≥2 个竖线（至少 3 列），覆盖 docx 表格/xlsx 行/Markdown 表格"""
    return line.count("|") >= 2


def _split_structural_blocks(text: str) -> list[str]:
    """把文本拆成结构块：连续表格行整体成块（不切断），其余按空行分段落"""
    blocks: list[str] = []
    para: list[str] = []
    table: list[str] = []

    def flush_para():
        if para:
            blocks.append("\n".join(para))
            para.clear()

    def flush_table():
        if table:
            blocks.append("\n".join(table))
            table.clear()

    for line in text.split("\n"):
        if _is_table_row(line):
            flush_para()
            table.append(line.strip())
        else:
            flush_table()
            if line.strip():
                para.append(line.strip())
            else:
                flush_para()
    flush_para()
    flush_table()
    return [b for b in blocks if b.strip()]


def _split_table_block(table: str) -> list[str]:
    """超大表格按行分组，每组重复首行（表头），保住列语义不丢"""
    rows = [r for r in table.split("\n") if r.strip()]
    if len(rows) <= 1:
        return [table]
    header = rows[0]
    size, chunks, group = config.CHUNK_SIZE, [], [header]
    used = len(header)
    for r in rows[1:]:
        if used + len(r) + 1 > size and len(group) > 1:
            chunks.append("\n".join(group))
            group, used = [header], len(header)
        group.append(r)
        used += len(r) + 1
    if len(group) > 1:
        chunks.append("\n".join(group))
    return chunks


def _chunk_structured(text: str) -> list[str]:
    """结构化装箱：结构块贪心合并至 CHUNK_SIZE；超大表格行分组、超大段落滑窗兜底"""
    size = config.CHUNK_SIZE
    chunks: list[str] = []
    cur = ""
    for block in _split_structural_blocks(text):
        if len(block) > size:
            if cur:
                chunks.append(cur)
                cur = ""
            if _is_table_row(block.split("\n")[0]):
                chunks.extend(_split_table_block(block))
            else:
                chunks.extend(_window_split(block))
        elif len(cur) + len(block) + 2 <= size:
            cur = f"{cur}\n\n{block}" if cur else block
        else:
            chunks.append(cur)
            cur = block
    if cur:
        chunks.append(cur)
    return [c.strip() for c in chunks if c.strip()]


def _breadcrumb_prefix(sec: str, stack: list[tuple[int, str]]) -> str:
    """小节切片的父标题链前缀（breadcrumb）：`###` 子节脱离 `#` 主标题后
    语义残缺（检索命中「配置步骤」不知道属于哪个产品/主题），伤害召回。
    返回形如 "[产品手册 > 部署]\\n" 的前缀；无父标题返回空串。
    预算控制在 120 字符内（超长标题链截断，不挤占正文空间）。"""
    m = re.match(r"(#{1,4})\s+(.+)", sec)
    level = len(m.group(1)) if m else 1
    parents = [t for lv, t in stack if lv < level]
    if not parents:
        return ""
    chain = " > ".join(parents)[-120:]
    return f"[{chain}]\n"


def chunk_text(text: str) -> list[str]:
    """结构化切片：Markdown 先按标题分段；各段内表格整体保留（超大表格行分组
    重复表头），段落贪心装箱；无结构文本直接结构化装箱（滑窗仅作超大段落兜底）。
    独立成片的子节携带父标题链前缀（breadcrumb），保住层级上下文。"""
    if _HEADING_RE.search(text):
        sections = [p.strip() for p in _HEADING_RE.split(text) if p.strip()]
        chunks: list[str] = []
        buffer = ""
        stack: list[tuple[int, str]] = []   # 标题栈：[(层级, 标题文本)]
        for sec in sections:
            if len(sec) > config.CHUNK_SIZE:
                if buffer:
                    chunks.append(buffer)
                    buffer = ""
                prefix = _breadcrumb_prefix(sec, stack)
                pieces = _chunk_structured(sec)
                if prefix:
                    # 首片带完整正文（含本节标题），后续片补父标题链防脱离上下文
                    chunks.extend(pieces[:1])
                    chunks.extend(f"{prefix}{p}" for p in pieces[1:])
                else:
                    chunks.extend(pieces)
            elif len(buffer) + len(sec) + 1 <= config.CHUNK_SIZE:
                buffer = f"{buffer}\n\n{sec}" if buffer else sec
            else:
                chunks.append(buffer)
                buffer = sec
            # 维护标题栈（本节标题入栈，供后续子节取父链）
            m = re.match(r"(#{1,4})\s+(.+)", sec)
            if m:
                level, title = len(m.group(1)), m.group(2).strip()
                stack = [(lv, t) for lv, t in stack if lv < level]
                stack.append((level, title))
        if buffer:
            chunks.append(buffer)
        return [c for c in chunks if c.strip()]
    return _chunk_structured(text)


def chunk_single_file(root: str, name: str) -> list[dict]:
    """加载并切片单个文件，返回 [{source, text, page?}]；解析失败返回 [] 不阻断。

    全量建库（load_chunks）与增量索引（VectorStore.rebuild_incremental）
    共用本函数，保证两条路径的切片行为完全一致。
    PDF 只解析一次：按页提取成功时正文直接由各页文本派生——
    原实现先跑整篇提取判空、再按页重解析一遍，每个 PDF 建库成本翻倍。
    """
    ext = os.path.splitext(name)[1].lower()
    path = os.path.join(root, name)
    chunks = []
    if ext == ".pdf":
        try:
            pages = _extract_pdf_pages(path)
        except Exception as e:  # noqa: BLE001
            logger.warning(f"PDF 按页切片失败，退回整篇切片 {name}: {e}")
            pages = None
        if pages is not None:
            for page_no, page_text in pages:
                page_text = (page_text or "")[:_MAX_EXTRACT_CHARS]
                if not page_text.strip():
                    continue
                for piece in chunk_text(page_text):
                    item = {"source": name, "text": piece}
                    if page_no:
                        item["page"] = page_no
                    chunks.append(item)
            return chunks
        # 按页提取失败：退回整篇提取（单次解析兜底），继续走通用路径
    try:
        if ext in _EXTRACTORS:
            text = _EXTRACTORS[ext](path)
        else:
            with open(path, encoding="utf-8") as f:
                text = f.read()
    except Exception as e:  # noqa: BLE001 - 坏文件不能阻断整个建库流程
        logger.warning(f"文档解析失败，已跳过 {name}: {e}")
        return []
    text = (text or "")[:_MAX_EXTRACT_CHARS]
    if not text.strip():
        return []
    for piece in chunk_text(text.strip()):
        chunks.append({"source": name, "text": piece})
    return chunks


def load_chunks(knowledge_dir: str | None = None) -> list[dict]:
    """加载并切片，返回 [{source, text, page?}]（text 为切片）。

    PDF 按页独立切片并携带 page 元数据（引用溯源/原文预览定位的地基）；
    其余格式不带页码。单文件失败不阻断（与 load_documents 策略一致）。
    """
    root = knowledge_dir or config.KNOWLEDGE_DIR
    result = []
    if not os.path.isdir(root):
        return result
    for name in sorted(os.listdir(root)):
        if os.path.splitext(name)[1].lower() not in SUPPORTED_EXTS:
            continue
        result.extend(chunk_single_file(root, name))
    return result
